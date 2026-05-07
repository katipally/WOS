from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

out_path = r"C:\Users\winston\WOS\DATA298B_WOS_Final_Report_DRAFT.docx"

doc = Document()

title = doc.add_paragraph("WOS — AI Agent Desktop Application: Final Project Report (Draft)")
title.runs[0].bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Author: Winston (wcc0) [fill in full name]").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Course: DATA 298B MSDA Project II").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f"Date: {date.today().isoformat()}").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

def h1(t):
    p = doc.add_paragraph(t)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

def h2(t):
    p = doc.add_paragraph(t)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

def add_known_unknown(known, unknown):
    if known:
        doc.add_paragraph("Known from repo:")
        for x in known:
            doc.add_paragraph(f"- {x}")
    if unknown:
        doc.add_paragraph("Missing / not evidenced in repo (needs input):")
        for x in unknown:
            doc.add_paragraph(f"- {x}")

h1("1. Introduction")

h2("1.1 Project Background and Executive Summary")
add_known_unknown(
    known=[
        "WOS is an Electron desktop application (Electron Forge + Vite) with a React UI.",
        "The app is an AI agent orchestrator that routes between model providers (OpenAI, Anthropic, Hugging Face Spaces) and executes tools via a structured tool-calling contract.",
        "The system prompt includes explicit policies: reuse known context, all clarifying questions via AskUser, and routing meeting/project work to subagents via Task.",
        "The project includes a synthetic data generation pipeline for orchestration fine-tuning and Colab notebooks for training.",
        "Hugging Face Spaces integration exists to serve OpenAI-compatible endpoints and use them as models inside WOS.",
    ],
    unknown=[
        "Specific external stakeholders / organizational context / deployment target users.",
        "Quantified motivation/problem statement and success metrics.",
    ],
)

doc.add_paragraph(
    "Planned approach (as implemented in repo): fine-tune an orchestration model using synthetic tool trajectories, "
    "validate/repair generated JSONL for strict schema adherence, and deploy the resulting model behind an OpenAI-compatible API (HF Space) used by the desktop app."
)

h2("1.2 Project Requirements")
add_known_unknown(
    known=[
        "Functional requirements evidenced by code: connect apps (Slack/GitHub/Jira/Google), run tool calls, ask user via AskUser, delegate to meeting/projects subagents via Task.",
        "AI requirements implied by training: tool-call correctness, recovery from injected tool failures, policy adherence.",
        "Data requirements: JSONL trajectories with strict roles/messages/tool_calls schema and metadata splits.",
    ],
    unknown=[
        "Explicit measurable acceptance criteria for each feature.",
        "Formal security/privacy requirements for connected apps and meeting data.",
    ],
)

h2("1.3 Project Deliverables")
add_known_unknown(
    known=[
        "Desktop application source code (Electron main + React renderer).",
        "Synthetic dataset generator and dataset audit/repair tooling.",
        "Training notebooks for orchestration, meeting, and coding models (Colab).",
        "HF Space deployment assets and quickstart docs.",
    ],
    unknown=[
        "Final packaged binaries and a demo video link.",
    ],
)

h2("1.4 Technology and Solution Survey")
doc.add_paragraph(
    "Insufficient evidence in repo to provide a complete survey of alternative commercial solutions and comparisons. "
    "Partial technologies evidenced by the codebase are listed below."
)
add_known_unknown(
    known=[
        "Model provider integration: OpenAI, Anthropic, and OpenAI-compatible servers via Hugging Face Spaces.",
        "Client: Electron + React + TypeScript; storage: SQLite (drizzle-orm); testing: Vitest + Playwright.",
    ],
    unknown=[
        "Systematic comparison against competing desktop agent apps (requires external research).",
    ],
)

h2("1.5 Literature Survey of Existing Research")
doc.add_paragraph(
    "Not enough information in the repo to cite specific research papers with APA references. This section requires external literature selection and citation."
)

h1("2. Data and Project Management Plan")

h2("2.1 Data Management Plan")
add_known_unknown(
    known=[
        "Orchestration training data stored as JSONL under training-data/, with generator scripts, audit, and repair utilities.",
        "Splits are handled via metadata and stable hashing logic.",
        "The runtime app stores conversations and settings in SQLite; API keys and app credentials are stored encrypted.",
    ],
    unknown=[
        "Formal data retention policy and compliance posture.",
        "Exact storage locations for production logs and whether user data is uploaded.",
    ],
)

h2("2.2 Project Development Methodology")
doc.add_paragraph(
    "Repo evidence indicates iterative development with lint/test scripts, E2E tests, and a synthetic-data-to-fine-tune pipeline. "
    "A formal lifecycle diagram is not included in the repo."
)

h2("2.3 Project Organization Plan")
doc.add_paragraph(
    "Not enough information in repo to produce a work breakdown structure (WBS) by team member."
)

h2("2.4 Project Resource Requirements and Plan")
add_known_unknown(
    known=[
        "Local dev requires Node.js + npm; Electron runtime; Python for dataset tooling; optional Colab GPU for fine-tuning.",
        "Inference hosting options include HF Spaces (Docker) and external providers.",
    ],
    unknown=[
        "Costed hardware plan (GPU tier costs, budgets) and license costs.",
    ],
)

h2("2.5 Project Schedule")
doc.add_paragraph("Not enough information in repo for a Gantt/PERT with task owners and dates.")

h1("3. Data Engineering")

h2("3.1 Data Process")
add_known_unknown(
    known=[
        "Raw and cleaned orchestration datasets exist; the generator produces new trajectories using a teacher model and validates them.",
        "Audit/repair scripts canonicalize system prompts and enforce strict message/tool schema.",
    ],
    unknown=[
        "Any real user data provenance (not present in repo).",
    ],
)

doc.add_paragraph("Sections 3.2–3.7 require dataset samples and visualizations not currently available in the repo snapshot.")

h1("4. Model Development")

h2("4.1 Model Proposals")
add_known_unknown(
    known=[
        "Orchestration model: Qwen3-32B LoRA fine-tune (adapter repo wcc0/wos_orch_qwen; base unsloth/Qwen3-32B-bnb-4bit).",
        "Meeting and coding notebooks propose QLoRA SFT strategies for specialized assistants.",
        "Intent classifier uses a fast external model (default claude-haiku-4-5-20251001) for tool filtering.",
    ],
    unknown=[
        "Final selected meeting/coding model repos and evaluation results.",
    ],
)

h2("4.2 Model Supports")
doc.add_paragraph("Provider adapters and tool execution run in the Electron main process. Fine-tuning is performed in Colab with Unsloth + TRL SFTTrainer.")

h2("4.3–4.5 Model Comparison / Evaluation")
doc.add_paragraph("Not enough empirical results in repo to provide comparisons, metrics, and validated outcomes. Insert after running evals.")

h1("5. Data Analytics and Intelligent System")

h2("5.1 System Requirements Analysis")
doc.add_paragraph("System boundary evidenced by code: desktop app that connects to external services (Slack/GitHub/Jira/Google), executes tools, and uses LLM providers.")

h2("5.2 System Design")
add_known_unknown(
    known=[
        "Architecture: Electron main process (providers/tools/db) + preload (IPC) + React renderer (UI).",
        "Data repository: SQLite tables for conversations, messages, settings, API keys, app connections, MCP servers.",
    ],
    unknown=[
        "Formal architecture diagrams and stable UI mockups.",
    ],
)

h2("5.3 Intelligent Solution")
doc.add_paragraph("Intelligence is delivered through (a) an orchestration model producing tool calls under policy constraints; (b) specialized subagents invoked via Task; (c) optional intent analysis to reduce tool exposure.")

h2("5.4 System Supporting Environment")
doc.add_paragraph("Technologies evidenced: Electron Forge + Vite, React 19, TypeScript, SQLite, OpenAI/Anthropic SDKs, Playwright/Vitest, Hugging Face Spaces deployment.")

h1("6. System Evaluation and Visualization")
doc.add_paragraph("Not enough in-repo evidence for the full evaluation/visualization sections. Add model execution metrics, runtime performance measurements, and UI screenshots after evaluation runs.")

h1("7. Conclusion")
doc.add_paragraph("Conclusion sections require final eval results and demonstration evidence. Fill after training and blind-test evaluation.")

h1("References")
doc.add_paragraph("References are not provided in repo. Add APA citations for models, libraries, and any research papers used.")

h1("Appendices")
doc.add_paragraph("Appendix placeholders: system testing screenshots, data storage links, code/demo artifacts.")

for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

doc.save(out_path)
print(out_path)
